import os
import logging
from docx import Document
from const import constants as c

# logging setup
logging.basicConfig(filename='merge_tables.log', level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')


def get_docx_files(directory):
    """
    Get a list of all docx files in the directory
    :param directory: path to the directory containing the docx files
    :return: docx files list
    """
    return [f for f in os.listdir(directory) if f.endswith('.docx')]


def merge_tables(directory, output_file):
    """
    Merges tables from all docx files in the given directory into a single table and saves it as a new docx file.
    Assumes that all tables have the same structure and that the first column of each table contains unique values
    :param directory: path to the directory containing the docx files
    :param output_file: result file
    :return: True if the table merge was successful, otherwise False
    """

    docx_files = get_docx_files(directory)

    # open the first file and get the first table
    try:
        doc = Document(os.path.join(directory, docx_files[0]))
        table = doc.tables[0]
    except Exception as e:
        logging.error(f'Error opening file {docx_files[0]}: {e}')
        return False

    # create a new document
    new_doc = Document()
    new_table = new_doc.add_table(rows=0, cols=len(table.columns))
    new_table.style = table.style

    # iterate over the remaining files and add rows from the tables
    try:
        first_col_unique_values = set()
        for filename in docx_files:
            doc = Document(os.path.join(directory, filename))
            for i, row in enumerate(doc.tables[0].rows):
                if row.cells[0].text.strip() not in first_col_unique_values:
                    first_col_unique_values.add(row.cells[0].text.strip())
                    new_row = new_table.add_row().cells
                    for j, cell in enumerate(row.cells):
                        new_row[j].text = cell.text
    except Exception as e:
        logging.error(f'Error processing file {filename}: {e}')
        return False

    # save the new document
    try:
        new_doc.save(os.path.join(directory, output_file))
    except Exception as e:
        logging.error(f'Error saving file {output_file}: {e}')
        return False

    return True


if __name__ == '__main__':
    if merge_tables(c.DIRECTORY, c.OUTPUT_FILE):
        logging.info("Tables merged successfully")
    else:
        logging.error("An error occurred during table merging")
